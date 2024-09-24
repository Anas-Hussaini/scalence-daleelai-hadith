from openai import OpenAI
import logging
from dotenv import load_dotenv
import os
import chromadb
import chromadb.utils.embedding_functions as embedding_functions
from instaloader import Instaloader, Post
import requests
from pydub import AudioSegment
import docx
import json
from config import (
    dotenv_path,
    vectorstore_path,
    collection_name,
    output_dir,
    username,
    transcription_model,
    transcription_language,
    llm_model,
    embedding_model
)


load_dotenv(dotenv_path=dotenv_path,override=True)
openai_token = os.environ['OpenAI_TOKEN']

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    filename='app.log',
    filemode='a'
)

# Also log to console
console = logging.StreamHandler()
console.setLevel(logging.INFO)
formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
console.setFormatter(formatter)
logging.getLogger('').addHandler(console)

def download_instagram_video_mp3(post_url, output_dir=output_dir):
    logging.info("Starting the download process...")
    
    # Ensure the output directory exists
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        logging.info(f"Created directory {output_dir}")
    
    try:
        # Create an instance of Instaloader
        L = Instaloader(download_videos=True)
        logging.info("Instaloader instance created.")
        
        # Load session file
        L.load_session_from_file(username=username) 
        logging.info("Session loaded successfully.")
        
        # Extract the shortcode from the URL
        shortcode = post_url.split("/")[-2]
        video_filename = f"{output_dir}/{shortcode}.mp4"
        audio_filename = f"{output_dir}/{shortcode}.mp3"

        # Skip download if the MP3 file already exists
        if os.path.exists(audio_filename):
            logging.info("Audio file already exists. Skipping download.")
            return audio_filename
        
        # Download the post using the shortcode
        post = Post.from_shortcode(L.context, shortcode)
        video_url = post.video_url
    
        # Send an HTTP GET request to the video URL
        response = requests.get(video_url)
        
        # Check if the request was successful
        if response.status_code == 200:
            # Save the video file
            with open(video_filename, 'wb') as file:
                for chunk in response.iter_content(chunk_size=1024):
                    if chunk:
                        file.write(chunk)
            logging.info(f'Video downloaded to {video_filename}')
            
            # Convert the video file to MP3
            logging.info("Converting video to MP3...")
            video = AudioSegment.from_file(video_filename, format="mp4")
            video.export(audio_filename, format="mp3")
            logging.info(f'Audio saved as {audio_filename}')
            
            # Optionally, delete the video file after conversion
            os.remove(video_filename)
            
            logging.info(f'Removed temporary video file {video_filename}')
            
            return audio_filename
        else:
            logging.error(f'Failed to download video. HTTP status code: {response.status_code}')
    except Exception as e:
        logging.error(f'An error occurred in downloading or converting the video: {e}', exc_info=True)
    
def transcript(audio_filename:str):
    try:
        # Initialize OpenAI client
        client = OpenAI(api_key=openai_token)
        logging.info("OpenAI client initialized successfully.")

        # Open the audio file
        audio_path = f"{output_dir}/{audio_filename}.mp3"
        with open(audio_path, "rb") as audio_file:
            logging.info(f"Opened audio file: {audio_path}")

            # Transcribe the audio file using Whisper
            repsonse = client.audio.transcriptions.create(
                model=transcription_model,
                file=audio_file,
                language=transcription_language
            )
            
            logging.info("Transcription completed successfully.")
            transcription = repsonse.text 
        return transcription
    except Exception as e:
        logging.error(f"Error during transcription: {e}", exc_info=True)

def llm_layer(transcription):
    try:
        # Initialize OpenAI client
        client = OpenAI(api_key=openai_token)
        logging.info("OpenAI client initialized successfully.")
        
        openai_response = client.chat.completions.create(
            model=llm_model,
            messages=[
                {"role": "system", "content": "You are an assistant tasked with extracting only the hadith content in english from the text provided by the user. Ignore any additional commentary, or unrelated text. Provide only the extracted hadith."},
                {"role": "user", "content": transcription}
            ]
            )

        hadith_crux =  openai_response.choices[0].message.content
            
        return hadith_crux
    
    except Exception as e:
        logging.error(f"Error in LLM layer processing: {e}", exc_info=True)

def query(hadith_crux:str):
    try:
        # Initialize ChromaDB client
        vectorstore_client = chromadb.PersistentClient(path=vectorstore_path)
        logging.info("ChromaDB client initialized successfully.")

        # Prepare the embedding function
        openai_ef = embedding_functions.OpenAIEmbeddingFunction(
            api_key=openai_token,
            model_name=embedding_model
        )
        logging.info("Embedding function prepared successfully.")

        # Set up the collection in ChromaDB
        collection = vectorstore_client.get_collection(
            name=collection_name,
            embedding_function=openai_ef
        )
        logging.info(f"ChromaDB collection '{collection_name}' set up successfully.")
        
        retrieved_docs_llm_layer = collection.query(query_texts=hadith_crux, n_results=3)
        logging.info("Query with llm layer executed successfully.")
        
        return retrieved_docs_llm_layer
    
    except Exception as e:
        logging.error(f"Error querying ChromaDB: {e}", exc_info=True)
        raise


def output_parse(retrieved_docs_llm_layer,transcription,hadith_crux):
    
    try:
        data_llm_layer = retrieved_docs_llm_layer
        metadatas_llm_layer = data_llm_layer['metadatas'][0]
        distances_llm_layer = data_llm_layer['distances'][0]
        documents_llm_layer = data_llm_layer['documents'][0]


        # Create a new Document object
        doc = docx.Document()

        # Add transcription section
        doc.add_heading('Transcription:', level=1)
        doc.add_paragraph(transcription)

        # Add a separator
        doc.add_paragraph('-' * 50)

        # Add Hadith Extraction section
        doc.add_heading('Hadith Extraction through llm layer:', level=1)
        doc.add_paragraph(hadith_crux)

        doc.add_paragraph('-' * 50)

        doc.add_heading('Results with LLM Layer:', level=1)
        
        # Add documents with metadata and distance
        for i in range(len(documents_llm_layer)):
            doc.add_heading(f'Document {i + 1}:', level=2)
            doc.add_paragraph(documents_llm_layer[i])
            
            # Add Metadata
            doc.add_heading('Metadata:', level=3)
            doc.add_paragraph(json.dumps(metadatas_llm_layer[i], indent=4))
            
            # Add Distance
            doc.add_heading('Distance:', level=3)
            doc.add_paragraph(str(distances_llm_layer[i]))
            
            # Add a separator for each document
            doc.add_paragraph('-' * 50)
            
            

        logging.info(f"Results saved successfully to .....")

    except Exception as e:
        logging.error(f"An error occurred: {e}", exc_info=True)

    return doc